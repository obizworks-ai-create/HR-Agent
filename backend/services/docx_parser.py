from docx import Document
import io

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Handles both standard paragraphs and tables (common in resumes).
    """
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []
        
        # 1. Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # 2. Extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                            
        return '\n'.join(full_text).strip()
        
    except Exception as e:
        print(f"Error extracting DOCX text with python-docx: {e}")
        return ""
